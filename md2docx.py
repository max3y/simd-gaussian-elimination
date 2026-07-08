#!/usr/bin/env python3
"""
Markdown → Word (.docx) 转换器
专用于将"SIMD向量化高斯消元优化实验报告"从 Markdown 转为
符合学术排版规范的 Word 文档。
"""

import re
import sys
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# =============================================================================
# 全局样式配置
# =============================================================================
TITLE_FONT_SIZE = Pt(18)
H1_FONT_SIZE = Pt(15)
H2_FONT_SIZE = Pt(13)
H3_FONT_SIZE = Pt(12)
BODY_FONT_SIZE = Pt(11)
CODE_FONT_SIZE = Pt(9)
TABLE_FONT_SIZE = Pt(9)
FONT_NAME = '宋体'
FONT_NAME_EN = 'Times New Roman'
CODE_FONT = 'Consolas'

BODY_LINE_SPACING = 1.5
CODE_LINE_SPACING = 1.15

def set_cell_border(cell, **kwargs):
    """设置表格单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            element = OxmlElement(f'w:{edge}')
            for attr in ['sz', 'val', 'color', 'space']:
                if attr in edge_data:
                    element.set(qn(f'w:{attr}'), str(edge_data[attr]))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_table_borders(table):
    """给整个表格添加边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)
    tblPr.append(borders)

def set_paragraph_spacing(paragraph, line_spacing=1.5, space_after=Pt(6)):
    """设置段落间距"""
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_after = space_after

def add_formatted_paragraph(doc, text, font_size=BODY_FONT_SIZE, bold=False,
                             alignment=None, font_name=FONT_NAME,
                             font_name_en=FONT_NAME_EN):
    """添加格式化段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = font_size
    run.bold = bold
    run.font.name = font_name_en
    r = run._element
    rFonts = r.find(qn('w:rPr'))
    if rFonts is None:
        rFonts = OxmlElement('w:rPr')
        r.insert(0, rFonts)
    eastAsia = rFonts.find(qn('w:rFonts'))
    if eastAsia is None:
        eastAsia = OxmlElement('w:rFonts')
        rFonts.append(eastAsia)
    eastAsia.set(qn('w:eastAsia'), font_name)
    if alignment is not None:
        p.alignment = alignment
    set_paragraph_spacing(p, BODY_LINE_SPACING)
    return p

def add_code_block(doc, code_text):
    """添加代码块(灰底+等宽字体)"""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        run = p.add_run(line if line else ' ')
        run.font.size = CODE_FONT_SIZE
        run.font.name = CODE_FONT
        set_paragraph_spacing(p, CODE_LINE_SPACING, Pt(0))
        # 灰色背景
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'F5F5F5')
        shading.set(qn('w:val'), 'clear')
        p._element.get_or_add_pPr().append(shading)

def add_table_from_data(doc, headers, rows):
    """从数据创建格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = TABLE_FONT_SIZE
        run.font.name = FONT_NAME_EN
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 表头灰底
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D9E2F3')
        shading.set(qn('w:val'), 'clear')
        cell._element.get_or_add_tcPr().append(shading)

    # 数据行
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[1 + r].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = TABLE_FONT_SIZE
            run.font.name = FONT_NAME_EN
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, 1.0, Pt(0))

    doc.add_paragraph()  # 表后间距
    return table

def add_section_heading(doc, text, level=1):
    """添加章节标题"""
    sizes = {1: H1_FONT_SIZE, 2: H2_FONT_SIZE, 3: H3_FONT_SIZE}
    size = sizes.get(level, BODY_FONT_SIZE)
    p = add_formatted_paragraph(doc, text, font_size=size, bold=True,
                                 alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_paragraph_spacing(p, BODY_LINE_SPACING, Pt(12))
    return p

def add_math_paragraph(doc, text):
    """添加数学公式段落(斜体+居中)"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = BODY_FONT_SIZE
    run.italic = True
    run.font.name = FONT_NAME_EN
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, BODY_LINE_SPACING, Pt(8))
    return p

# =============================================================================
# 主转换函数
# =============================================================================
def convert_md_to_docx(md_path, docx_path):
    doc = Document()

    # ---- 页面设置 ----
    section = doc.sections[0]
    section.page_width = Cm(21.0)   # A4
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_buffer = []
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # --- 代码块处理 ---
        if line.strip().startswith('```'):
            if in_code_block:
                add_code_block(doc, '\n'.join(code_buffer))
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # --- 空行 ---
        if not line.strip():
            in_table = False
            i += 1
            continue

        # --- 标题 # / ## / ### ---
        if line.startswith('# ') and not line.startswith('## '):
            title = line[2:].strip()
            # 跳过一级标题中的"基于SIMD..."使其成为文档主标题
            p = doc.add_paragraph()
            run = p.add_run(title)
            run.font.size = TITLE_FONT_SIZE
            run.bold = True
            run.font.name = FONT_NAME_EN
            r = run._element
            rFonts = r.find(qn('w:rPr'))
            if rFonts is None:
                rFonts = OxmlElement('w:rPr')
                r.insert(0, rFonts)
            eastAsia = OxmlElement('w:rFonts')
            eastAsia.set(qn('w:eastAsia'), FONT_NAME)
            rFonts.append(eastAsia)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, BODY_LINE_SPACING, Pt(16))
            i += 1
            continue

        if line.startswith('## '):
            add_section_heading(doc, line[3:].strip(), level=2)
            i += 1
            continue

        if line.startswith('### '):
            add_section_heading(doc, line[4:].strip(), level=3)
            i += 1
            continue

        # --- 水平线 ---
        if line.strip() == '---':
            p = doc.add_paragraph()
            pPr = p._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # --- 表格 ---
        if line.strip().startswith('|') and line.strip().endswith('|'):
            # 跳过分隔行
            if re.match(r'^\|[\s\-:|]+\|$', line.strip()):
                i += 1
                continue
            row = [cell.strip() for cell in line.strip().split('|')[1:-1]]
            table_rows.append(row)
            in_table = True
            # 检查下一行是否还是表格
            if i + 1 < len(lines) and lines[i + 1].strip().startswith('|'):
                i += 1
                continue
            else:
                # 表格结束，输出
                if table_rows:
                    headers = table_rows[0]
                    data = table_rows[1:]
                    # 判断是否是真正的表格(表头不含**)
                    if all(not h.startswith('**') for h in headers):
                        add_table_from_data(doc, headers, data)
                    else:
                        # 简化的表格说明文字 — 作为正文
                        for row_data in table_rows:
                            add_formatted_paragraph(doc,
                                '  '.join(row_data).replace('**', ''))
                table_rows = []
                in_table = False
                i += 1
                continue

        # --- 数学公式 $$...$$ ---
        if line.strip().startswith('$$') and line.strip().endswith('$$'):
            math_text = line.strip()[2:-2].strip()
            add_math_paragraph(doc, math_text)
            i += 1
            continue

        # --- 引用块 ---
        if line.startswith('>'):
            quote_text = line[1:].strip()
            # 合并连续的多行引用
            while i + 1 < len(lines) and lines[i + 1].startswith('>'):
                i += 1
                quote_text += ' ' + lines[i][1:].strip()
            p = doc.add_paragraph()
            run = p.add_run(quote_text)
            run.font.size = Pt(10)
            run.italic = True
            run.font.name = FONT_NAME_EN
            run.font.color.rgb = RGBColor(100, 100, 100)
            p.paragraph_format.left_indent = Cm(1.0)
            set_paragraph_spacing(p, BODY_LINE_SPACING)
            i += 1
            continue

        # --- 粗体标记 **text** ---
        # --- 普通正文 ---
        # 处理内联格式
        clean_line = line.strip()

        # 跳过空行
        if not clean_line:
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        # 处理 **粗体**
        segments = re.split(r'(\*\*[^*]+\*\*)', clean_line)
        for seg in segments:
            if seg.startswith('**') and seg.endswith('**'):
                run = p.add_run(seg[2:-2])
                run.bold = True
            else:
                # 处理行内代码 `code`
                code_segments = re.split(r'(`[^`]+`)', seg)
                for cs in code_segments:
                    if cs.startswith('`') and cs.endswith('`'):
                        run = p.add_run(cs[1:-1])
                        run.font.name = CODE_FONT
                        run.font.size = CODE_FONT_SIZE
                    else:
                        run = p.add_run(cs)
            run.font.size = BODY_FONT_SIZE
            run.font.name = FONT_NAME_EN
            r = run._element
            rFonts = r.find(qn('w:rPr'))
            if rFonts is None:
                rFonts = OxmlElement('w:rPr')
                r.insert(0, rFonts)
            eastAsia = rFonts.find(qn('w:rFonts'))
            if eastAsia is None:
                eastAsia = OxmlElement('w:rFonts')
                rFonts.append(eastAsia)
            eastAsia.set(qn('w:eastAsia'), FONT_NAME)

        set_paragraph_spacing(p, BODY_LINE_SPACING)
        i += 1

    # ---- 保存 ----
    doc.save(docx_path)
    print(f"[DONE] Word document saved to: {docx_path}")

# =============================================================================
if __name__ == '__main__':
    md_file = r'd:\桌面\parallel\rewritten\实验报告_SIMD向量化高斯消元优化.md'
    docx_file = r'd:\桌面\parallel\rewritten\实验报告_SIMD向量化高斯消元优化.docx'
    convert_md_to_docx(md_file, docx_file)
